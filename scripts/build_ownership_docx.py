from docx import Document
from docx.shared import Pt

doc = Document()
normal = doc.styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(11)

def H(text, level=1): doc.add_heading(text, level=level)
def P(text, bold=False, italic=False):
    p = doc.add_paragraph(); r = p.add_run(text); r.bold = bold; r.italic = italic; return p
def B(text): doc.add_paragraph(text, style='List Bullet')
def CODE(text):
    p = doc.add_paragraph(); r = p.add_run(text); r.font.name='Consolas'; r.font.size=Pt(8.5); return p
def TABLE(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers)); t.style='Light Grid Accent 1'
    hdr = t.rows[0].cells
    for i,h in enumerate(headers):
        hdr[i].text=h
        for par in hdr[i].paragraphs:
            for run in par.runs: run.bold=True
    for row in rows:
        cells=t.add_row().cells
        for i,v in enumerate(row): cells[i].text=v
    return t

doc.add_heading('Project Skeleton — Ownership Map (Corrected — Full Product)', level=0)
P('Frontend / Backend / AI-Services structure, team rotation for whole-project experience, and the phase-wise plan', italic=True)
P('Stack: React (frontend SPA) + FastAPI (backend API) + Python AI-services (RAG + model serving: vLLM/Ollama, embeddings, OpenAI router) + PostgreSQL + Vector DB + Docker/K8s.')
P('Tenancy: Multi-university multi-tenant SaaS. Auth: SSO (SAML/OIDC) + email-password JWT.')
P('Ownership philosophy: Collective ownership with rotation. No file is permanently locked to one squad. Each pillar has a rotating Primary (builds) and Reviewer (reviews + learns); pods rotate across pillars each phase so every team member gets hands-on experience across the whole project.', italic=True)

H('1. Monorepo Folder Structure (Frontend / Backend / AI)', 1)
tree = """project-root/
│
├── frontend/                       # React SPA - Student / Professor / Admin dashboards
│   ├── public/
│   ├── src/
│   │   ├── app/                    # routing, providers, layout
│   │   ├── auth/                   # login, SSO callback, JWT/session
│   │   ├── modules/
│   │   │   ├── student/            # AI chat, my-attendance, my-marks
│   │   │   ├── professor/          # attendance, marks, analytics
│   │   │   ├── admin/              # payments, payroll, halltickets, expenses
│   │   │   └── tenant-admin/       # tenant CRUD, enable/disable AI, branding
│   │   ├── components/ui/          # shared design system
│   │   ├── lib/                    # API client, tenant resolver, hooks
│   │   └── store/
│   ├── package.json
│   └── .env.example
│
├── backend/                        # FastAPI - API + business logic (NO templates)
│   ├── app/
│   │   ├── main.py                 # bootstrap + router registration
│   │   ├── config.py
│   │   ├── database.py
│   │   ├── core/
│   │   │   ├── tenancy/            # tenant_id context, RLS helpers
│   │   │   └── security/           # JWT, RBAC, SSO
│   │   ├── models/                 # tenant, session, user, audit_log, academic, finance
│   │   ├── schemas/                # Pydantic request/response
│   │   ├── routes/                 # tenants, launch, auth, academic, finance, health
│   │   ├── services/               # tenant_service, launch_service, token_service, ...
│   │   └── clients/
│   │       └── ai_client.py        # talks to ai-services over HTTP (decoupled)
│   ├── migrations/                 # Alembic
│   ├── tests/
│   ├── requirements.txt
│   └── .env.example
│
├── ai-services/                    # AI / RAG layer (scales on GPUs, deploys separately)
│   ├── rag/
│   │   ├── ingestion/              # parse -> chunk -> embed professor notes
│   │   ├── retrieval/              # vector search + re-rank (tenant/course scoped)
│   │   ├── pipeline/               # build grounded prompt -> answer + citations
│   │   └── guardrails/             # grounding, moderation, anti-injection
│   ├── serving/
│   │   ├── llm_server/             # vLLM / Ollama config to serve self-hosted model
│   │   ├── embeddings/             # embedding model service (bge / e5)
│   │   └── router/                 # self-hosted <-> OpenAI routing logic
│   ├── prompts/                    # versioned prompt templates
│   ├── model_registry/             # CONFIG ONLY - model names, versions, endpoints
│   │   └── models.yaml             # WEIGHTS live in object storage, NOT here
│   ├── workers/                    # async embedding / reindex jobs
│   ├── tests/
│   ├── requirements.txt
│   └── .env.example
│
├── deploy/                         # IaC & deployment (shared)
│   ├── docker/                     # frontend.Dockerfile, backend.Dockerfile, ai.Dockerfile
│   ├── k8s/                        # frontend / backend / ai-services / vector-db / ingress
│   └── terraform/                  # VPC, Postgres, GPU node pool, object storage, vector DB
│
├── docs/                           # architecture, data model, API ref, RAG pipeline
├── scripts/                        # seed, SIS/CSV import, model-download, maintenance
├── docker-compose.yml              # local dev: frontend + backend + ai + db + vector-db
├── .gitignore                      # ignores model weights, node_modules, .env
└── README.md"""
CODE(tree)

H('2. Key Architectural Rules', 1)
B('Backend never imports AI code directly - it calls ai-services over HTTP via backend/app/clients/ai_client.py, keeping the GPU-heavy AI layer independently deployable.')
B('Model weights are NOT in the repo - model_registry/models.yaml lists names/versions/endpoints only; binaries are pulled from object storage / model registry at container startup (scripts/model-download). .gitignore blocks weight files.')
B('Vector DB and object storage are infrastructure, provisioned in deploy/terraform/ - not folders in the codebase.')
B('Every backend table carries tenant_id with PostgreSQL Row-Level Security; vector collections and storage prefixes are tenant-scoped.')
B('Server-rendered Jinja templates/static are removed - all UI is handled by the React frontend/.')

H('3. Pillar Build Plan', 1)
TABLE(
    ['Pillar', 'Key Areas', 'What Gets Built'],
    [
        ['Frontend (React)', 'auth, student, professor, admin, tenant-admin, ui, lib', 'Role-aware dashboards, AI chat UI, API client, SSO + JWT flows'],
        ['Backend (FastAPI)', 'core (tenancy/security), models, schemas, routes, services, clients', 'Multi-tenant API, auth/RBAC, tenant & academic & finance logic, ai_client'],
        ['AI-Services (Python)', 'rag (ingestion/retrieval/pipeline/guardrails), serving, router, workers', 'Note ingestion + embeddings, vector search, grounded answers, model routing'],
        ['Infra & QA', 'deploy (docker/k8s/terraform), tests, migrations, scripts', 'Containers, cluster manifests, IaC, CI tests (incl. cross-tenant isolation)'],
    ]
)

H('4. Cross-Functional Ownership Model (everyone learns the whole project)', 1)
P('Principle: collective ownership + rotation across the three pillars. Every developer rotates through Frontend, Backend, AI-Services, and Infra/QA so that by the end each member has built, reviewed, and tested all three pillars.')

H('4.1 Roles per pillar (rotating)', 2)
B('Primary (P): builds the pillar work this phase.')
B('Reviewer (R): reviews PRs, pairs, and learns - becomes next phase Primary.')
B('Lead: owns the API contracts between pillars (frontend<->backend<->ai), architecture, and standards.')

H('4.2 Rotation matrix (Primary focus per phase)', 2)
P('4 pods (2-3 devs each) shift one pillar to the right every phase, so each pod touches Frontend, Backend, AI-Services, and Infra/QA.')
TABLE(
    ['Phase', 'Pod 1', 'Pod 2', 'Pod 3', 'Pod 4'],
    [
        ['Phase 1 - Foundation', 'Frontend shell + auth UI', 'Backend core + tenancy/security', 'AI-services scaffold + serving', 'Infra: Docker, K8s, Terraform, CI'],
        ['Phase 2 - Core', 'Backend tenant/academic API', 'AI-services RAG ingestion + retrieval', 'Infra: vector DB, model registry, tests', 'Frontend tenant-admin + dashboards'],
        ['Phase 3 - AI Integration', 'AI-services pipeline + guardrails + router', 'Infra/tests + isolation suite', 'Frontend student AI chat + professor analytics', 'Backend ai_client + finance API'],
        ['Phase 4 - Hardening', 'Infra: scaling, DR, observability', 'Frontend polish + accessibility', 'Backend security/audit hardening', 'AI-services cost routing + eval'],
    ]
)
P('After 4 phases every pod has worked across all three pillars plus infra. Cross-pod pairing is required on integration days.', italic=True)

H('4.3 Guardrails so rotation does not cause chaos', 2)
B('Stable API contracts (OpenAPI for backend, typed client for frontend, defined ai-services HTTP interface) so pillars stay decoupled while people rotate.')
B('One coding standard per pillar (ESLint/Prettier for frontend; ruff/black for Python) enforced in CI.')
B('Mandatory PR review by the incoming Primary - that is how knowledge transfers.')
B('Lead reviews any change to cross-pillar contracts, requirements/package files, and IaC.')
B('Daily 15-min sync where each pod demos shipped work; weekly cross-pillar knowledge share.')

H('5. Phase-Wise Work Breakdown', 1)

H('Phase 1 - Foundation', 2)
P('Goal: all three pillars boot and talk to each other (empty).', bold=True)
for b in ['Frontend: app shell, routing, auth screens (SSO + JWT), API client stub.',
          'Backend: config, database, core tenancy + security, main.py with health route.',
          'AI-services: service scaffold, serving config (vLLM/Ollama), embeddings stub, models.yaml.',
          'Infra: Dockerfiles, docker-compose (frontend+backend+ai+db+vector-db), CI pipeline, base Terraform.']:
    B(b)
P('Exit criteria: docker-compose up runs all three pillars; frontend reaches backend; backend reaches ai-services health.', italic=True)

H('Phase 2 - Core Modules', 2)
P('Goal: tenant + academic + ingestion work in isolation.', bold=True)
for b in ['Backend: tenant CRUD + enable/disable-AI, auth/RBAC, academic (attendance/marks) APIs, audit log.',
          'AI-services: note ingestion (parse->chunk->embed) and tenant/course-scoped vector retrieval.',
          'Frontend: tenant-admin module + role dashboards reading real APIs.',
          'Infra: provision vector DB, model registry/weights download, set up automated tests.']:
    B(b)
P('Exit criteria: a tenant can be created and configured; professor notes ingest into the vector store; dashboards render real data.', italic=True)

H('Phase 3 - AI Integration', 2)
P('Goal: end-to-end grounded AI experience.', bold=True)
for b in ['AI-services: RAG pipeline (grounded prompt -> answer + citations), guardrails, self-hosted<->OpenAI router.',
          'Backend: ai_client wiring, AI launch/session + JWT, finance APIs.',
          'Frontend: student AI chat with citations; professor analytics + weak-area views.',
          'Infra/QA: cross-tenant isolation test suite, full-flow integration tests.']:
    B(b)
P('Exit criteria: student asks a question -> grounded answer with citations; full create->enable->launch->chat flow passes tests; no cross-tenant leakage.', italic=True)

H('Phase 4 - Hardening & Scale', 2)
P('Goal: production-ready, multi-university scale.', bold=True)
for b in ['Infra: autoscaling, GPU node pool, DR/backups, observability (metrics/logs/traces).',
          'Backend: security/audit hardening, rate limits, per-tenant quotas.',
          'AI-services: cost-aware routing (self-hosted-first), answer quality eval harness.',
          'Frontend: accessibility, responsive polish, error/empty states.',
          'Docs + knowledge share: every member can explain all three pillars.']:
    B(b)
P('Exit criteria: all tests green; observability live; cost controls in place; team fully cross-trained.', italic=True)

H('6. Quick Reference - Pillars (no permanent owners)', 1)
TABLE(
    ['Pillar', 'Folders', 'Notes'],
    [
        ['Frontend', 'frontend/src/{app,auth,modules,components,lib,store}', 'Rotates each phase'],
        ['Backend', 'backend/app/{core,models,schemas,routes,services,clients} + migrations', 'Rotates each phase'],
        ['AI-Services', 'ai-services/{rag,serving,prompts,model_registry,workers}', 'Rotates each phase'],
        ['Infra & QA', 'deploy/{docker,k8s,terraform}, tests/, scripts/', 'Rotates each phase'],
        ['Architecture (Lead)', 'cross-pillar API contracts, docker-compose.yml, README.md, docs/', 'Stable lead role'],
    ]
)

P('This version restructures the project into Frontend / Backend / AI-Services pillars, keeps weights out of the repo via a model registry, decouples backend from AI via an HTTP client, and remaps the team rotation and phase plan across the three pillars so every member gains whole-project experience.', italic=True)

out='/app/scripts/Project_Skeleton_Ownership_Map_CORRECTED.docx'
doc.save(out); print('Saved:', out)
