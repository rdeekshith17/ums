from docx import Document
from docx.shared import Pt

doc = Document()
normal = doc.styles['Normal']; normal.font.name='Calibri'; normal.font.size=Pt(11)

def H(t,l=1): doc.add_heading(t,level=l)
def P(t,bold=False,italic=False):
    p=doc.add_paragraph(); r=p.add_run(t); r.bold=bold; r.italic=italic; return p
def B(t): doc.add_paragraph(t, style='List Bullet')
def CHK(t): doc.add_paragraph('[ ]  '+t, style='List Bullet')
def CODE(t):
    p=doc.add_paragraph(); r=p.add_run(t); r.font.name='Consolas'; r.font.size=Pt(8.5); return p
def TABLE(headers, rows, style='Light Grid Accent 1'):
    t=doc.add_table(rows=1, cols=len(headers)); t.style=style
    h=t.rows[0].cells
    for i,x in enumerate(headers):
        h[i].text=x
        for par in h[i].paragraphs:
            for run in par.runs: run.bold=True
    for row in rows:
        c=t.add_row().cells
        for i,v in enumerate(row): c[i].text=v
    return t

doc.add_heading('Phase 1 — Foundation: Detailed Work & Workflow (per Pod)', level=0)
P('Multi-tenant University AI SaaS — Frontend / Backend / AI-Services / Infra', italic=True)
P('Goal: Every pillar (Frontend, Backend, AI-Services) boots, is containerized, and can talk to its neighbor with no real features yet. By end of phase, "docker-compose up" brings the whole system online and a health request flows Frontend -> Backend -> AI-Services -> DB/Vector-DB.', bold=True)
P('Suggested duration: 5 working days (1 sprint).')

H('0. Pod Composition & Focus This Phase', 1)
TABLE(['Pod','Phase-1 Focus','Devs'],
[['Pod 1','Frontend shell + auth UI','2-3'],
 ['Pod 2','Backend core (tenancy + security)','2-3'],
 ['Pod 3','AI-Services scaffold + serving','2-3'],
 ['Pod 4','Infra: Docker, K8s, Terraform, CI','2-3'],
 ['Lead','Cross-pillar API contracts, standards, integration','1']])

H('1. Dependency & Sequence Map', 1)
CODE("""Day 0 (Lead) --> Repo + monorepo skeleton + API contracts agreed
        |
   +----+--------------------+--------------------+--------------------+
   v                         v                    v                    v
Pod 4 (Infra)          Pod 2 (Backend)      Pod 3 (AI-svc)       Pod 1 (Frontend)
docker-compose+CI      config/db/core       scaffold+health      shell+auth UI
(everyone depends      (needs DB from       (no external dep)    (needs backend
 on this FIRST)         Pod 4 compose)                            health contract)
   |                         |                    |                    |
   +--------> INTEGRATION CHECKPOINT (Day 4): health call traverses all pillars <--+
""")
P('Critical path: Lead contracts (Day 0-1) -> Pod 4 docker-compose (Day 1) -> all other pods build against it. Pod 4 MUST deliver compose + shared .env.example by end of Day 1 or other pods are blocked.', italic=True)

H('2. Lead — Day 0/1 (must finish first)', 1)
B('Create monorepo: empty frontend/ backend/ ai-services/ deploy/ docs/ scripts/ + README, .gitignore (ignore node_modules, .env, model weights), docker-compose.yml placeholder.')
B('Branch & PR strategy: main (protected) -> develop -> feat/<pod>-<task>. PRs need 1 review by incoming Primary + green CI.')
B('Agree 3 API contracts in docs/contracts/: Frontend<->Backend (base URL, /api/health shape, Authorization Bearer, tenant resolution); Backend<->AI (internal URL, /health + /v1/echo); standard response envelope {data, error, request_id}.')
B('Coding standards: ESLint/Prettier (frontend), ruff/black (Python) wired into CI.')
B('Post Definition of Done checklist for the whole team.')
P('Lead deliverable: repo scaffold + docs/contracts/*.md + DoD merged to develop by end of Day 1.', italic=True)

H('3. Pod 4 — Infrastructure (highest priority, unblocks others)', 1)
TABLE(['Day','Work','Deliverable'],
[['1','docker-compose.yml with 5 services: frontend, backend, ai-services, postgres, vector-db (pgvector/Qdrant); shared network + .env.example','docker-compose up boots DB + Vector DB; placeholder app containers'],
 ['1-2','3 Dockerfiles in deploy/docker/ (frontend, backend, ai) - multi-stage dev+prod','Each pillar builds into an image'],
 ['2-3','CI pipeline (lint+build+test stub) on every PR to develop','Green CI; PRs blocked if red'],
 ['3-4','Base Terraform in deploy/terraform/ (VPC, Postgres, object storage, vector DB, GPU node-pool) - plan only','terraform plan runs clean'],
 ['4-5','Base K8s manifests in deploy/k8s/ (namespaces, deployments, services, ingress) - not deployed','kubectl --dry-run validates']])
P('DoD: any dev runs docker-compose up and gets all 5 services healthy; CI runs on PRs; Terraform/K8s skeletons reviewed.', italic=True)

H('4. Pod 2 — Backend Core (tenancy + security)', 1)
TABLE(['Day','Work','Deliverable'],
[['1','config.py (settings from .env) + database.py (SQLAlchemy engine + session)','App reads env; DB session opens against compose Postgres'],
 ['1-2','core/tenancy/ middleware resolves tenant_id (subdomain/header) into request context + scoping helper','Every request has tenant_id (stub tenant)'],
 ['2','core/security/ JWT verify dependency + RBAC role skeleton (no real login)','Protected route rejects missing/invalid token'],
 ['2-3','main.py bootstrap + register routers + /api/health (checks DB)','GET /api/health returns agreed envelope'],
 ['3','clients/ai_client.py thin HTTP client calling AI-services /health + /v1/echo','Backend reaches AI-services'],
 ['3-4','Alembic setup + empty initial migration (no business tables)','alembic upgrade head runs in compose'],
 ['4','Wire /api/health to also report AI-services reachability','Health OK only when DB + AI-svc reachable']])
P('DoD: /api/health returns OK, validates a JWT on a protected stub, resolves tenant context, confirms AI-services connectivity.', italic=True)

H('5. Pod 3 — AI-Services Scaffold + Serving', 1)
TABLE(['Day','Work','Deliverable'],
[['1','Scaffold ai-services app with /health','Service boots in compose'],
 ['1-2','serving/llm_server/ config for self-hosted model (vLLM/Ollama) in mock/CPU mode (no GPU dependency)','LLM serve config present; mock mode works on CPU'],
 ['2','serving/embeddings/ stub + model_registry/models.yaml (names/versions/endpoints; NO weights in repo)','Registry validated; scripts/model-download placeholder'],
 ['2-3','serving/router/ skeleton (self-hosted vs OpenAI interface; returns mock)','Router interface callable'],
 ['3','/v1/echo stub endpoint per Backend<->AI contract','Backend ai_client gets valid response'],
 ['3-4','Folder scaffolds: rag/ingestion, rag/retrieval, rag/pipeline, rag/guardrails, workers/ (empty + READMEs)','Structure ready for Phase 2']])
P('DoD: AI-services boots in compose, /health + /v1/echo respond per contract, model registry present, weights excluded from git.', italic=True)
P('FLAG TO MANAGER: To avoid GPU dependency in Phase 1, the LLM runs in MOCK/CPU mode. Real model serving is deferred to Phase 2/3 once GPU infra (Pod 4 Terraform) is provisioned.', bold=True)

H('6. Pod 1 — Frontend Shell + Auth UI', 1)
TABLE(['Day','Work','Deliverable'],
[['1','Bootstrap React app (app/ routing, providers, layout); env-based backend URL','App runs in compose, renders shell'],
 ['1-2','lib/ API client (reads backend URL, attaches Authorization header, parses envelope)','Client wired to backend'],
 ['2','auth/ login screen + SSO button (UI only) + JWT/session storage scaffolding','Login page renders; token storage stubbed'],
 ['2-3','Role-aware routing: empty student/, professor/, admin/, tenant-admin/ shells','Navigating by role shows correct shell'],
 ['3','System status page calling /api/health showing backend + AI-svc status','Frontend -> Backend -> AI health visible in UI'],
 ['3-4','components/ui/ base design-system primitives + theme','Shared UI kit ready for Phase 2']])
P('DoD: frontend boots in compose, shows role-based shell, status page proves full chain Frontend->Backend->AI reachable.', italic=True)

H('7. Day-by-Day Gantt (X = active, * = key event)', 1)
TABLE(['Workstream','D1','D2','D3','D4','D5'],
[['Lead: repo + contracts','X','','','','*'],
 ['Pod 4: docker-compose + .env','X','X','','',''],
 ['Pod 4: Dockerfiles + CI','X','X','X','',''],
 ['Pod 4: Terraform + K8s','','','X','X','X'],
 ['Pod 2: config/db/core','X','X','','',''],
 ['Pod 2: security + health + ai_client','','X','X','X',''],
 ['Pod 2: Alembic migration','','','X','X',''],
 ['Pod 3: scaffold + health','X','X','','',''],
 ['Pod 3: serving + registry + router','','X','X','X',''],
 ['Pod 3: /v1/echo + rag scaffolds','','','X','X',''],
 ['Pod 1: app shell + API client','X','X','','',''],
 ['Pod 1: auth UI + role shells','','X','X','',''],
 ['Pod 1: status page + UI kit','','','X','X',''],
 ['INTEGRATION CHECKPOINT','','','','*',''],
 ['Hardening + Demo + Retro','','','','','*']])

H('8. Daily Workflow', 1)
B('Day 0-1 Setup & contracts: Lead lands repo + contracts; Pod 4 lands compose + .env.example. Other pods scaffold against compose stubs.')
B('Day 2 Parallel build: each pod builds its skeleton against agreed contracts; mock the neighbor if not ready. No pod waits on another internal code, only on the contract.')
B('Day 3 First wiring: Backend<->AI /v1/echo; Frontend<->Backend /api/health. Open integration PRs.')
B('Day 4 Integration checkpoint (key event): whole team runs docker-compose up; verify one health request traverses Frontend -> Backend -> AI -> DB/Vector-DB. Fix breakages together.')
B('Day 5 Hardening + demo: green CI, README updated, each pod demos; Lead confirms exit criteria; retro + plan Phase-2 rotation.')
P('Rituals: 15-min daily standup (shipped/blocked/next); contract changes go through the Lead; PRs reviewed by the incoming Primary of that pillar (knowledge rotates into next phase).')

H('9. Phase-1 Exit Criteria (Demo Gate Checklist)', 1)
CHK('docker-compose up boots all 5 services healthy.')
CHK('Frontend renders role-based shells + login/SSO UI (no real auth yet).')
CHK('GET /api/health returns OK, validates JWT on a protected stub, resolves tenant_id.')
CHK('Backend reaches AI-services (/v1/echo); AI-services serves in mock mode.')
CHK('A single health request visibly traverses Frontend -> Backend -> AI -> DB/Vector-DB.')
CHK('CI (lint + build + test stub) green on develop; Terraform plan and K8s --dry-run validate.')
CHK('docs/contracts/ complete; README updated; every pod has demoed.')

H('10. Decisions to Confirm Before Kickoff', 1)
B('Vector DB choice for compose: pgvector (simplest) vs Qdrant (scales better)?')
B('Confirm LLM mock mode in Phase 1 is acceptable (real GPU serving deferred to Phase 2)?')

out='/app/scripts/Project_Phase1_Detailed_Plan.docx'
doc.save(out); print('Saved:', out)
