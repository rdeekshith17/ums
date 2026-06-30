from docx import Document
from docx.shared import Pt

doc = Document()
n = doc.styles['Normal']; n.font.name='Calibri'; n.font.size=Pt(11)

def H(t,l=1): doc.add_heading(t,level=l)
def P(t,b=False,i=False):
    p=doc.add_paragraph(); r=p.add_run(t); r.bold=b; r.italic=i; return p
def BUL(t): doc.add_paragraph(t, style='List Bullet')
def CODE(t):
    p=doc.add_paragraph(); r=p.add_run(t); r.font.name='Consolas'; r.font.size=Pt(8.5)
def TB(headers, rows):
    t=doc.add_table(rows=1, cols=len(headers)); t.style='Light Grid Accent 1'
    h=t.rows[0].cells
    for i,x in enumerate(headers):
        h[i].text=x
        for par in h[i].paragraphs:
            for rr in par.runs: rr.bold=True
    for row in rows:
        c=t.add_row().cells
        for i,v in enumerate(row): c[i].text=v

doc.add_heading('UniAI — Admin Dashboard: Detailed Documentation & File-Level Coding Spec', level=0)
P('Multi-tenant University AI Dashboard — first deliverable screen', i=True)
P('Stack: React frontend + FastAPI backend + PostgreSQL. Data already modelled in backend/app/models (17 tables) and seeded for 8 Telangana universities.')

# ---------------------------------------------------------------- 1
H('1. Purpose & Scope', 1)
P('The first screen to ship is the Admin Dashboard. The admin (Registrar / Accounts / Exam Controller) sees, scoped to their own university (tenant_id): KPIs, students, professors & payroll, fees/payments, hall tickets, expenses, departments/courses, and the audit log.')
P('This document specifies every dashboard field and the exact table.column it comes from, the backend API surface that feeds it, and which file each pod/teammate builds.', b=True)

# ---------------------------------------------------------------- 2
H('2. Admin Dashboard — Layout', 1)
P('Recommended layout (top to bottom): KPI cards row -> tabbed sections (Students | Professors & Payroll | Payments | Hall Tickets | Expenses | Departments | Audit Log).')

H('2.1 KPI Cards (Overview row)', 2)
TB(['Card / Field','Source (table.column)','Computation'],
[['Total Students','student_profiles','COUNT where tenant_id'],
 ['Total Professors','professor_profiles','COUNT where tenant_id'],
 ['Total Admins','admin_profiles','COUNT where tenant_id'],
 ['Departments','departments','COUNT where tenant_id'],
 ['Courses','courses','COUNT where tenant_id'],
 ['Fees Collected','payments.amount','SUM where status=success'],
 ['Fees Pending','invoices.amount','SUM of unpaid+overdue + 50% of partial'],
 ['Total Expenses','expenses.amount','SUM where tenant_id'],
 ['Monthly Payroll','payslips.net','SUM for latest payroll_run'],
 ['AI Status','universities.ai_enabled','Boolean badge (Enabled/Disabled)']])

H('2.2 Students Table', 2)
TB(['Column','Source'],
[['Roll No','student_profiles.roll_no'],
 ['Name','users.full_name'],
 ['Email','users.email'],
 ['Department','departments.name'],
 ['Batch / Year','student_profiles.batch_year / current_year'],
 ['CGPA','student_profiles.cgpa'],
 ['Attendance %','attendance_records (present/total per student)'],
 ['Fee Status','invoices.status'],
 ['Hall Ticket','hall_tickets.eligibility']])

H('2.3 Professors & Payroll', 2)
TB(['Column','Source'],
[['Employee ID','professor_profiles.employee_id'],
 ['Name','users.full_name'],
 ['Designation','professor_profiles.designation'],
 ['Department','departments.name'],
 ['Specialization','professor_profiles.specialization'],
 ['Monthly Salary','professor_profiles.monthly_salary'],
 ['Latest Net Pay','payslips.net (latest run)'],
 ['Deductions','payslips.deductions']])

H('2.4 Payments / Invoices', 2)
TB(['Column','Source'],
[['Invoice ID','invoices.id'],
 ['Student','users.full_name via student_profiles'],
 ['Description','invoices.description'],
 ['Amount','invoices.amount'],
 ['Status','invoices.status'],
 ['Due Date','invoices.due_date'],
 ['Paid Amount','payments.amount'],
 ['Method','payments.method'],
 ['Txn Ref','payments.provider_ref']])

H('2.5 Hall Tickets', 2)
TB(['Column','Source'],
[['Student','users.full_name via student_profiles'],
 ['Roll No','student_profiles.roll_no'],
 ['Exam','hall_tickets.exam_name'],
 ['Eligibility','hall_tickets.eligibility'],
 ['Block Reason','hall_tickets.block_reason'],
 ['Issued At','hall_tickets.issued_at']])

H('2.6 Expenses', 2)
TB(['Column / Widget','Source'],
[['Category','expenses.category'],
 ['Description','expenses.description'],
 ['Amount','expenses.amount'],
 ['Vendor','expenses.vendor'],
 ['Spent On','expenses.spent_on'],
 ['Expense-by-Category chart','GROUP BY expenses.category (SUM amount)']])

H('2.7 Departments & Courses', 2)
TB(['Column','Source'],
[['Department','departments.name / code'],
 ['# Students','student_profiles COUNT by department_id'],
 ['# Courses','courses COUNT by department_id'],
 ['HOD','professor_profiles where designation=HOD']])

H('2.8 Audit Log', 2)
TB(['Column','Source'],
[['Event Type','audit_logs.event_type'],
 ['Details','audit_logs.details'],
 ['Tenant','audit_logs.tenant_id'],
 ['Timestamp','audit_logs.created_at']])

# ---------------------------------------------------------------- 3
H('3. Backend API Surface (feeds the dashboard)', 1)
P('All endpoints are tenant-scoped: the admin JWT carries tenant_id and every query filters by it. All under /api prefix.')
TB(['Method & Path','Returns','Used by'],
[['GET /api/admin/overview','All KPI card values','KPI cards'],
 ['GET /api/admin/students','Paginated students + attendance% + fee/hallticket status','Students tab'],
 ['GET /api/admin/professors','Professors + latest payslip net','Professors & Payroll tab'],
 ['GET /api/admin/payments','Invoices + payment info','Payments tab'],
 ['GET /api/admin/hall-tickets','Hall ticket list + eligibility','Hall Tickets tab'],
 ['GET /api/admin/expenses','Expenses + category summary','Expenses tab'],
 ['GET /api/admin/departments','Departments with counts + HOD','Departments tab'],
 ['GET /api/admin/audit-logs','Recent audit events','Audit Log tab']])

# ---------------------------------------------------------------- 4
H('4. File-Level Coding Spec (who builds what)', 1)
P('The Admin Dashboard mainly involves Pod 2 (Backend) and Pod 1 (Frontend). Pod 4 (Infra/QA) owns seed/tests. Pod 3 (AI) is not involved in this screen.')

H('4.1 Pod 2 — Backend (FastAPI)', 2)
TB(['File','Status','What it must contain'],
[['app/database.py','DONE','Engine, SessionLocal, Base, get_db.'],
 ['app/models/*.py','DONE','17 ORM models (university, user, academics, finance, audit).'],
 ['app/core/tenancy.py','BUILD','Dependency that extracts tenant_id from JWT/subdomain and yields it to every admin query.'],
 ['app/core/security.py','BUILD','JWT verify + require_role("admin") dependency.'],
 ['app/schemas/admin.py','BUILD','Pydantic response models: OverviewKPIs, StudentRow, ProfessorRow, PaymentRow, HallTicketRow, ExpenseRow, DepartmentRow, AuditRow.'],
 ['app/services/admin_service.py','BUILD','All aggregation/query logic: get_overview(), list_students() (joins users+dept, computes attendance %), list_professors() (joins latest payslip), list_payments(), list_hall_tickets(), list_expenses()+category summary, list_departments() counts, recent_audit(). Every query filters by tenant_id.'],
 ['app/routes/admin.py','BUILD','The 8 GET endpoints in section 3; each depends on get_current_admin + tenant_id, calls admin_service, returns the schema.'],
 ['app/main.py','BUILD','Register admin router: app.include_router(admin.router, prefix="/api/admin").']])

H('4.2 Pod 1 — Frontend (React)', 2)
TB(['File','What it must contain'],
[['src/lib/api.js','Axios/fetch client: base URL from env, attaches Authorization Bearer token, parses {data,error}.'],
 ['src/modules/admin/AdminDashboard.jsx','Page shell + tab navigation; fetches /api/admin/overview for KPI cards.'],
 ['src/modules/admin/components/KpiCards.jsx','Renders the 10 KPI cards (section 2.1).'],
 ['src/modules/admin/StudentsTable.jsx','Fetches /api/admin/students; table with columns from 2.2 + search/pagination.'],
 ['src/modules/admin/ProfessorsPayroll.jsx','Fetches /api/admin/professors; columns from 2.3.'],
 ['src/modules/admin/Payments.jsx','Fetches /api/admin/payments; columns from 2.4 + status badges.'],
 ['src/modules/admin/HallTickets.jsx','Fetches /api/admin/hall-tickets; eligible/blocked badges (2.5).'],
 ['src/modules/admin/Expenses.jsx','Fetches /api/admin/expenses; table + category bar chart (2.6).'],
 ['src/modules/admin/Departments.jsx','Fetches /api/admin/departments; counts + HOD (2.7).'],
 ['src/modules/admin/AuditLog.jsx','Fetches /api/admin/audit-logs; event timeline (2.8).'],
 ['src/components/ui/*','Shared Card, Table, Badge, Tabs, Chart primitives.']])

H('4.3 Pod 4 — Infrastructure & QA', 2)
TB(['File','What it must contain'],
[['backend/seed_data.py','DONE — sample data for all tables (used by the dashboard).'],
 ['backend/migrations/ (Alembic)','BUILD — versioned schema for the 17 tables instead of create_all.'],
 ['backend/tests/test_admin.py','BUILD — assert each /api/admin/* returns correct tenant-scoped data and rejects non-admins.'],
 ['deploy/docker-compose.yml','BUILD — frontend + backend + postgres wired for local run.']])

H('4.4 Pod 3 — AI-Services', 2)
P('Not involved in the Admin Dashboard. Engaged later for the Student RAG module.')

# ---------------------------------------------------------------- 5
H('5. Data Flow (one KPI request)', 1)
CODE("""Admin opens dashboard
  -> Frontend AdminDashboard.jsx calls GET /api/admin/overview (Bearer JWT)
  -> Backend security.py verifies JWT + role=admin
  -> tenancy.py extracts tenant_id from the token
  -> routes/admin.py -> admin_service.get_overview(db, tenant_id)
  -> service runs COUNT/SUM queries filtered by tenant_id
  -> returns OverviewKPIs schema (JSON)
  -> KpiCards.jsx renders the 10 cards""")

# ---------------------------------------------------------------- 6
H('6. Build Order (recommended)', 1)
BUL('Pod 2: security.py + tenancy.py (auth scaffolding) — unblocks every endpoint.')
BUL('Pod 2: schemas/admin.py -> admin_service.get_overview -> routes /api/admin/overview.')
BUL('Pod 1: api.js -> AdminDashboard.jsx + KpiCards.jsx (first visible result).')
BUL('Then iterate per tab: students -> professors/payroll -> payments -> hall tickets -> expenses -> departments -> audit, backend service + frontend table together.')
BUL('Pod 4: write test_admin.py alongside each endpoint; verify tenant isolation (admin of OU cannot see JNTUH data).')

P('Acceptance: Admin logs in -> sees KPI cards with real numbers for their university -> each tab loads tenant-scoped data -> a non-admin token is rejected with 403.', b=True)

out='/app/scripts/Admin_Dashboard_Documentation.docx'
doc.save(out); print('Saved:', out)
